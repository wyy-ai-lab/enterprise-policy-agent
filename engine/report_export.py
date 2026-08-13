"""
报告导出模块
支持 Markdown、Word、PDF 三种格式
"""

import io
import json
import os
import re
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple

from .dashboard import get_deadline_status, get_top_gaps, compute_dashboard_metrics
from .radar_chart import calculate_dimension_scores


# 诊断类型展示颜色（Word/PDF 复用）
DIAGNOSIS_COLORS = {
    "立即申报": {"hex": "#16a34a", "rgb": (22, 163, 74)},
    "培育申报": {"hex": "#d97706", "rgb": (217, 119, 6)},
    "持续关注": {"hex": "#2563eb", "rgb": (37, 99, 235)},
    "暂不适合": {"hex": "#dc2626", "rgb": (220, 38, 38)},
}

DIAGNOSIS_PRIORITY = {
    "立即申报": 0,
    "培育申报": 1,
    "持续关注": 2,
    "暂不适合": 3,
}

POLICY_PRIORITY = {
    "高": 0,
    "中": 1,
    "低": 2,
}


def _find_cjk_font() -> str:
    """查找中文字体路径：优先项目内置字体，再查找 Windows 系统字体"""
    # 1. 优先使用项目内置字体（兼容 Streamlit Cloud / Linux / macOS）
    bundled = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "fonts", "NotoSansSC-Regular.ttf")
    if os.path.exists(bundled):
        return bundled

    # 2.  fallback 到 Windows 常见字体目录
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return ""


def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """将 #rrggbb 转为 (r, g, b)"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def _avg_combined_score(results: List[Dict[str, Any]]) -> float:
    """计算平均综合匹配度"""
    scores = [r.get('combined_score', r.get('match_score', 0)) for r in results]
    return round(sum(scores) / len(scores), 1) if scores else 0.0


def select_top3_policies(results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    按「诊断类型优先级 → 综合分数降序 → 政策优先级 → 截止日由近到远」选出 TOP3 推荐政策

    返回：每条政策附带 timeline_advice、key_gaps、deadline_status 等展示字段
    """
    def _sort_key(r):
        diag_priority = DIAGNOSIS_PRIORITY.get(r.get('diagnosis', ''), 99)
        score = -(r.get('combined_score', r.get('match_score', 0)))
        policy_priority = POLICY_PRIORITY.get(r.get('priority', '中'), 1)
        status = get_deadline_status(r.get('deadline', ''))
        days = status['days_left'] if status['days_left'] is not None and status['days_left'] >= 0 else 9999
        return (diag_priority, score, policy_priority, days)

    sorted_results = sorted(results, key=_sort_key)
    top3 = sorted_results[:3]

    selected = []
    for rank, r in enumerate(top3, start=1):
        status = get_deadline_status(r.get('deadline', ''))
        diagnosis = r.get('diagnosis', '')

        if diagnosis == '立即申报':
            if status['is_urgent']:
                timeline = '建议本周内启动申报，避免错过截止日'
            else:
                timeline = '建议尽快启动申报，准备材料并提交'
        elif diagnosis == '培育申报':
            timeline = '建议 1-2 年内补齐差距并完成申报'
        elif diagnosis == '持续关注':
            timeline = '持续跟踪政策动态，待条件成熟后申报'
        else:
            timeline = '短期内差距较大，暂不推荐申报'

        key_gaps = []
        if r.get('failed'):
            key_gaps.extend([item.split('：')[0] if '：' in item else item.split(':')[0] for item in r['failed'][:3]])
        if r.get('unknown'):
            key_gaps.extend([item.split('：')[0] if '：' in item else item.split(':')[0] for item in r['unknown'][:2]])
        key_gaps = list(dict.fromkeys(key_gaps))[:3]  # 去重并限制数量

        selected.append({
            "rank": rank,
            "policy_id": r.get('policy_id', ''),
            "policy_name": r.get('policy_name', ''),
            "diagnosis": diagnosis,
            "priority": r.get('priority', '中'),
            "combined_score": r.get('combined_score', r.get('match_score', 0)),
            "deadline": r.get('deadline', ''),
            "deadline_status": status,
            "benefit": r.get('benefit', ''),
            "timeline_advice": timeline,
            "key_gaps": key_gaps,
            "reason": r.get('reason', ''),
        })

    return selected


def build_executive_summary(
    result: Dict[str, Any],
    capability_scores: Dict[str, int],
    top3_policies: List[Dict[str, Any]]
) -> str:
    """
    生成执行摘要文本，包含诊断概览、核心发现、TOP3 推荐和下一步行动建议
    """
    enterprise_name = result.get('enterprise_name', '未命名企业')
    results = result.get('results', [])
    summary = result.get('summary', {})
    total = len(results)
    avg_score = _avg_combined_score(results)

    is_enhanced = 'llm_config' in result
    mode_text = "硬条件 + LLM 软条件综合诊断" if is_enhanced else "硬条件诊断"

    # 最匹配政策：立即申报或培育申报中综合分数最高的
    actionable = [r for r in results if r.get('diagnosis') in ['立即申报', '培育申报']]
    best_policy = None
    if actionable:
        best_policy = max(actionable, key=lambda x: x.get('combined_score', x.get('match_score', 0)))

    # 最紧迫截止
    nearest = None
    nearest_days = None
    for r in results:
        status = get_deadline_status(r.get('deadline', ''))
        if status['days_left'] is not None and status['days_left'] >= 0:
            if nearest_days is None or status['days_left'] < nearest_days:
                nearest_days = status['days_left']
                nearest = r

    # 短板维度
    weak_dimensions = [dim for dim, score in capability_scores.items() if score < 60]

    # 高频缺失数据
    top_gaps = get_top_gaps(results, top_n=3)

    lines = []
    lines.append(f"本报告针对 **{enterprise_name}** 的政策申报机会进行了{mode_text}，共匹配 {total} 条政策。")
    lines.append(f"整体平均综合匹配度为 **{avg_score} 分**，其中：")
    lines.append(
        f"- 立即申报 **{summary.get('立即申报', 0)}** 条、培育申报 **{summary.get('培育申报', 0)}** 条、"
        f"持续关注 **{summary.get('持续关注', 0)}** 条、暂不适合 **{summary.get('暂不适合', 0)}** 条。"
    )
    lines.append("")
    lines.append("### 核心发现")

    if best_policy:
        lines.append(
            f"- **最匹配政策**：《{best_policy['policy_name']}》"
            f"（{best_policy['diagnosis']}，综合 {best_policy.get('combined_score', best_policy.get('match_score', 0))} 分）"
        )
    else:
        lines.append("- **最匹配政策**：当前无可立即申报或培育申报的政策，建议优先补齐基础数据。")

    if nearest:
        status = get_deadline_status(nearest.get('deadline', ''))
        lines.append(
            f"- **最紧迫截止**：《{nearest['policy_name']}》将于 {nearest.get('deadline', '')} 截止（{status['status_text']}）。"
        )
    else:
        lines.append("- **最紧迫截止**：暂无有效截止日信息。")

    if weak_dimensions:
        lines.append(f"- **短板维度**：{'、'.join(weak_dimensions)}，建议针对性提升。")
    else:
        lines.append("- **短板维度**：各能力维度均达到基本水平，保持并持续提升。")

    if top_gaps:
        gap_text = '；'.join([f"{name}（{count} 项）" for name, count in top_gaps])
        lines.append(f"- **高频缺失/差距项**：{gap_text}。")
    else:
        lines.append("- **高频缺失/差距项**：未发现显著高频差距项。")

    lines.append("")
    lines.append("### TOP3 推荐政策")
    if top3_policies:
        for p in top3_policies:
            lines.append(
                f"{p['rank']}. **{p['policy_name']}**（{p['diagnosis']}，{p['combined_score']} 分）"
                f" — {p['timeline_advice']}"
            )
    else:
        lines.append("暂无推荐政策。")

    lines.append("")
    lines.append("### 下一步行动建议")
    immediate = [r for r in results if r.get('diagnosis') == '立即申报']
    cultivate = [r for r in results if r.get('diagnosis') == '培育申报']

    if immediate:
        lines.append(f"- **立即行动（0-30 天）**：启动 {len(immediate)} 条「立即申报」政策的材料准备，优先处理临近截止的项目。")
    else:
        lines.append("- **立即行动（0-30 天）**：当前无立即申报政策，重点补充研发费用、知识产权等关键数据。")

    if cultivate:
        lines.append(f"- **中期培育（1-12 个月）**：针对 {len(cultivate)} 条「培育申报」政策制定补齐计划。")
    else:
        lines.append("- **中期培育（1-12 个月）**：持续积累资质和研发投入，为未来申报做准备。")

    lines.append("- **长期关注（1-2 年）**：建立政策监控机制，定期更新企业画像并重新诊断。")

    return "\n".join(lines)


def _build_enterprise_profile(enterprise: Dict[str, Any]) -> Dict[str, Any]:
    """
    整理企业画像快照，供报告展示使用
    """
    if not enterprise:
        return {"available": False, "text": "暂无企业画像数据。"}

    profile = {
        "available": True,
        "name": enterprise.get("name", "未命名企业"),
        "industry": enterprise.get("industry", "—"),
        "sub_industry": enterprise.get("sub_industry", "—"),
        "region": enterprise.get("region", "—"),
        "scale": enterprise.get("scale", "—"),
        "employees": enterprise.get("employees", "—"),
        "founded_year": enterprise.get("founded_year", "—"),
        "revenue": enterprise.get("revenue", "—"),
        "profit": enterprise.get("profit", "—"),
        "rd_investment": enterprise.get("rd_investment") or "—",
        "rd_ratio": enterprise.get("rd_ratio") or "—",
        "rd_team_size": enterprise.get("rd_team_size") or "—",
        "rd_team_ratio": enterprise.get("rd_team_ratio") or "—",
        "high_tech_income_ratio": enterprise.get("high_tech_income_ratio") or "—",
        "invention_patents": enterprise.get("invention_patents", 0),
        "utility_models": enterprise.get("utility_models", 0),
        "software_copyrights": enterprise.get("software_copyrights", 0),
        "trademarks": enterprise.get("trademarks", 0),
        "qualifications": enterprise.get("qualifications", []),
        "is_high_tech_enterprise": enterprise.get("is_high_tech_enterprise", False),
        "rd_accounting_system": enterprise.get("rd_accounting_system", False),
        "has_major_accident": enterprise.get("has_major_accident", False),
    }

    profile["text"] = (
        f"**{profile['name']}** 是一家位于 **{profile['region']}** 的 **{profile['scale']}**，"
        f"成立于 **{profile['founded_year']}** 年，所属行业为 **{profile['industry']}**（{profile['sub_industry']}）。"
        f"现有员工 **{profile['employees']}** 人，上年度营收 **{profile['revenue']}** 万元，利润 **{profile['profit']}** 万元。"
    )
    return profile


def build_report_sections(result: Dict[str, Any], capability_scores: Dict[str, int]) -> Dict[str, Any]:
    """
    统一生成报告各章节内容字典，供 Markdown / Word / PDF 复用
    """
    enterprise_name = result.get('enterprise_name', '未命名企业')
    diagnosis_date = result.get('diagnosis_date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    results = result.get('results', [])
    summary = result.get('summary', {})
    is_enhanced = 'llm_config' in result
    llm_config = result.get('llm_config', {})

    # 读取企业画像
    enterprise_file = "data/enterprise.json"
    enterprise: Dict[str, Any] = {}
    if os.path.exists(enterprise_file):
        with open(enterprise_file, 'r', encoding='utf-8') as f:
            enterprise = json.load(f)

    top3 = select_top3_policies(results)
    executive_summary = build_executive_summary(result, capability_scores, top3)
    avg_score = _avg_combined_score(results)
    metrics = compute_dashboard_metrics(result)
    gap_analysis = get_top_gaps(results, top_n=5)
    enterprise_profile = _build_enterprise_profile(enterprise)

    return {
        "title": "企业政策诊断报告",
        "enterprise_name": enterprise_name,
        "diagnosis_date": diagnosis_date,
        "mode": "硬条件 + LLM 软条件综合诊断" if is_enhanced else "硬条件诊断",
        "llm_provider": "演示模式" if llm_config.get('use_demo') else llm_config.get('provider', '未知'),
        "soft_score_count": llm_config.get('max_policies_for_soft_score', 0) if is_enhanced else 0,
        "total_policies": len(results),
        "summary_counts": summary,
        "avg_score": avg_score,
        "metrics": metrics,
        "enterprise_profile": enterprise_profile,
        "gap_analysis": gap_analysis,
        "executive_summary": executive_summary,
        "top3_policies": top3,
        "capability_scores": capability_scores,
        "capability_assessment": _build_capability_assessment(capability_scores),
        "detailed_results": results,
        "is_enhanced": is_enhanced,
        "next_steps": _build_next_steps(results),
    }


def _build_capability_assessment(capability_scores: Dict[str, int]) -> List[str]:
    """根据维度分数生成评价文本"""
    suggestions = []
    if capability_scores.get("经营规模", 0) < 60:
        suggestions.append("经营规模偏弱，可关注对中小微企业友好的专项政策。")
    if capability_scores.get("研发投入", 0) < 60:
        suggestions.append("研发投入不足，建议完善研发费用辅助账并提高研发投入占比。")
    if capability_scores.get("知识产权", 0) < 60:
        suggestions.append("知识产权较弱，建议申请发明专利、软件著作权等核心知识产权。")
    if capability_scores.get("资质荣誉", 0) < 60:
        suggestions.append("资质荣誉较少，建议申报高新技术企业、专精特新等资质。")
    if capability_scores.get("高新技术产业化", 0) < 60:
        suggestions.append("高新技术产品收入占比或市场证明不足，需补充相关证明材料。")
    if capability_scores.get("合规与成长", 0) < 60:
        suggestions.append("合规或成长数据存在短板，需排除安全质量事故并积累经营年限。")
    if not suggestions:
        suggestions.append("企业综合能力较为均衡，可重点争取高分值政策。")
    return suggestions


def _build_next_steps(results: List[Dict[str, Any]]) -> Dict[str, List[str]]:
    """生成下一步行动建议"""
    immediate = []
    medium = []
    long_term = []

    for r in results:
        if r.get('diagnosis') == '立即申报':
            immediate.append(f"启动《{r['policy_name']}》申报材料准备")
        elif r.get('diagnosis') == '培育申报':
            medium.append(f"制定《{r['policy_name']}》差距补齐计划")

    if not immediate:
        immediate.append("补充研发费用、知识产权等关键数据，争取更多立即申报机会")
    if not medium:
        medium.append("持续积累资质和研发投入，为培育申报做准备")

    long_term.extend([
        "建立政策监控机制，定期跟踪申报截止日和新政策",
        "每季度更新企业画像并重新诊断"
    ])

    return {
        "immediate": immediate,
        "medium": medium,
        "long": long_term,
    }


def build_markdown_report(result: Dict[str, Any]) -> str:
    """
    根据诊断结果生成 Markdown 报告
    """
    # 从企业画像读取能力分数；若诊断结果中已包含也可优先使用
    enterprise_file = "data/enterprise.json"
    if os.path.exists(enterprise_file):
        with open(enterprise_file, 'r', encoding='utf-8') as f:
            enterprise = json.load(f)
        capability_scores = enterprise.get('capability_scores') or calculate_dimension_scores(enterprise)
    else:
        capability_scores = calculate_dimension_scores({})

    sections = build_report_sections(result, capability_scores)

    report = f"""# {sections['title']}

**企业名称**：{sections['enterprise_name']}
**诊断时间**：{sections['diagnosis_date']}
**匹配政策数**：{sections['total_policies']} 条
**诊断模式**：{sections['mode']}
**平均综合匹配度**：{sections['avg_score']} 分
"""

    if sections['is_enhanced']:
        report += f"**LLM 提供商**：{sections['llm_provider']}  \n"
        report += f"**软条件评估政策数**：{sections['soft_score_count']} 条  \n"

    report += f"""
## 执行摘要

{sections['executive_summary']}

## 企业画像快照

"""
    ep = sections['enterprise_profile']
    if ep['available']:
        report += ep['text'] + "\n\n"
        report += "| 项目 | 内容 |\n|------|------|\n"
        report += f"| 所属行业 | {ep['industry']}（{ep['sub_industry']}） |\n"
        report += f"| 企业规模 | {ep['scale']}，员工 {ep['employees']} 人 |\n"
        report += f"| 经营数据 | 营收 {ep['revenue']} 万元 / 利润 {ep['profit']} 万元 |\n"
        report += f"| 研发投入 | {ep['rd_investment']} 万元，占比 {ep['rd_ratio']} |\n"
        report += f"| 研发团队 | {ep['rd_team_size']} 人，占比 {ep['rd_team_ratio']} |\n"
        report += f"| 高新技术产品收入占比 | {ep['high_tech_income_ratio']} |\n"
        report += f"| 知识产权 | 发明 {ep['invention_patents']} / 实用新型 {ep['utility_models']} / 软著 {ep['software_copyrights']} / 商标 {ep['trademarks']} |\n"
        report += f"| 已获资质 | {'、'.join(ep['qualifications']) if ep['qualifications'] else '—'} |\n"
        report += f"| 国家高新技术企业 | {'是' if ep['is_high_tech_enterprise'] else '否'} |\n"
        report += f"| 研发准备金制度 | {'已建立' if ep['rd_accounting_system'] else '未建立'} |\n"
        report += f"| 近三年重大事故 | {'有' if ep['has_major_accident'] else '无'} |\n"
    else:
        report += ep['text'] + "\n"

    report += """
## 诊断结果总览

| 诊断结果 | 数量 |
|---------|------|
"""
    for diagnosis in ["立即申报", "培育申报", "持续关注", "暂不适合"]:
        report += f"| {diagnosis} | {sections['summary_counts'].get(diagnosis, 0)} 条 |\n"

    metrics = sections['metrics']
    report += f"""
**平均综合匹配度**：{sections['avg_score']} 分
**紧急截止政策数**：{metrics['urgent_count']} 条
**已过期政策数**：{metrics['expired_count']} 条
"""
    if metrics['nearest_deadline']:
        report += f"**最近截止日**：{metrics['nearest_deadline']}（剩 {metrics['nearest_days']} 天）\n"

    report += """
## 高频差距与短板分析

"""
    if sections['gap_analysis']:
        report += "| 差距/缺失项 | 涉及政策数 |\n|-------------|-----------|\n"
        for name, count in sections['gap_analysis']:
            report += f"| {name} | {count} 条 |\n"
    else:
        report += "未发现显著高频差距项。\n"

    report += """
## TOP3 推荐政策路线图

"""
    if sections['top3_policies']:
        for p in sections['top3_policies']:
            deadline_text = f"截止：{p['deadline']}（{p['deadline_status']['status_text']}）" if p['deadline'] else "无固定截止日"
            report += f"""### {p['rank']}. {p['policy_name']}
- **诊断结果**：{p['diagnosis']}
- **综合分数**：{p['combined_score']} 分
- **政策优先级**：{p['priority']}
- **申报时间线**：{p['timeline_advice']}
- **{deadline_text}**
- **扶持内容**：{p['benefit']}
- **关键差距**：{'、'.join(p['key_gaps']) if p['key_gaps'] else '无'}

"""
    else:
        report += "暂无推荐政策。\n\n"

    report += """## 企业能力雷达图与维度分数

"""
    report += "| 能力维度 | 分数 | 评价 |\n"
    report += "|---------|------|------|\n"
    for dim, score in sections['capability_scores'].items():
        level = "强" if score >= 80 else "良" if score >= 60 else "弱"
        report += f"| {dim} | {score} 分 | {level} |\n"

    report += "\n### 维度短板建议\n\n"
    for item in sections['capability_assessment']:
        report += f"- {item}\n"

    report += """

## 全部政策诊断明细

| 诊断结果 | 数量 |
|---------|------|
"""
    for diagnosis in ["立即申报", "培育申报", "持续关注", "暂不适合"]:
        report += f"| {diagnosis} | {sections['summary_counts'].get(diagnosis, 0)} 条 |\n"

    report += """

## 详细诊断结果

"""

    diagnosis_order = ["立即申报", "培育申报", "持续关注", "暂不适合"]
    for diagnosis in diagnosis_order:
        items = [r for r in sections['detailed_results'] if r['diagnosis'] == diagnosis]
        if not items:
            continue

        report += f"### {diagnosis}（{len(items)} 条）\n\n"
        for r in items:
            report += f"#### {r['policy_name']}\n"

            if sections['is_enhanced'] and 'combined_score' in r:
                report += f"- **综合匹配度**：{r['combined_score']} 分（硬 {r['hard_score']} + 软 {r.get('soft_score', 'N/A')}）\n"
            else:
                report += f"- **匹配度**：{r['match_score']} 分\n"

            report += f"- **政策层级**：{r['level']}\n"
            report += f"- **申报截止**：{r['deadline']}\n"
            report += f"- **扶持内容**：{r['benefit']}\n"
            report += f"- **诊断理由**：{r['reason']}\n"

            if r['failed']:
                report += "- **差距**：\n"
                for item in r['failed']:
                    report += f"  - {item}\n"

            if r['unknown']:
                report += "- **需补充数据**：\n"
                for item in r['unknown']:
                    report += f"  - {item}\n"

            if sections['is_enhanced'] and r.get('soft_score') is not None:
                report += f"- **LLM 软条件评估**：{r['soft_score']} 分（置信度：{r.get('confidence', '未知')}）\n"
                report += f"- **综合评估**：{r.get('soft_assessment', '')}\n"

                if r.get('strengths'):
                    report += "- **优势**：\n"
                    for item in r['strengths']:
                        report += f"  - {item}\n"

                if r.get('weaknesses'):
                    report += "- **短板**：\n"
                    for item in r['weaknesses']:
                        report += f"  - {item}\n"

                if r.get('cultivation_suggestions'):
                    report += "- **培育建议**：\n"
                    for item in r['cultivation_suggestions']:
                        report += f"  - {item}\n"

            # 申报材料大纲（如果存在）
            if r.get('material_outline'):
                outline = r['material_outline']
                report += f"- **申报可行性**：{outline.get('applicability', '')}\n"
                report += "- **申报材料大纲**：\n"
                for section in outline.get('outline', []):
                    report += f"  - **{section.get('section', '')}**：{'; '.join(section.get('content', []))}\n"
                if outline.get('key_attachments'):
                    report += "- **关键附件清单**：\n"
                    for item in outline['key_attachments']:
                        report += f"  - {item}\n"
                if outline.get('gap_fill_plan'):
                    report += "- **差距补齐计划**：\n"
                    for item in outline['gap_fill_plan']:
                        report += f"  - {item}\n"
                if outline.get('notes'):
                    report += f"- **特别提醒**：{outline['notes']}\n"

            report += "\n"

    report += """## 下一步行动建议

### 立即行动（0-30 天）
"""
    for item in sections['next_steps']['immediate']:
        report += f"- {item}\n"

    report += "\n### 中期培育（1-12 个月）\n"
    for item in sections['next_steps']['medium']:
        report += f"- {item}\n"

    report += "\n### 长期关注（1-2 年）\n"
    for item in sections['next_steps']['long']:
        report += f"- {item}\n"

    report += """

---

*本报告由企业政策诊断辅导智能体自动生成*
"""
    return report


def build_word_report(
    result: Dict[str, Any],
    capability_scores: Optional[Dict[str, int]] = None,
    radar_image_bytes: Optional[bytes] = None
) -> bytes:
    """
    生成 Word 报告，返回 bytes

    参数：
        result: 诊断结果字典
        capability_scores: 企业能力维度分数（为 None 时从 data/enterprise.json 读取）
        radar_image_bytes: 雷达图图片 bytes（为 None 时使用文字表格 fallback）
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml

    if capability_scores is None:
        enterprise_file = "data/enterprise.json"
        if os.path.exists(enterprise_file):
            with open(enterprise_file, 'r', encoding='utf-8') as f:
                enterprise = json.load(f)
            capability_scores = enterprise.get('capability_scores') or calculate_dimension_scores(enterprise)
        else:
            capability_scores = calculate_dimension_scores({})

    sections = build_report_sections(result, capability_scores)
    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(sections['title'])
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 41, 55)

    doc.add_paragraph()
    cover_info = [
        ("企业名称", sections['enterprise_name']),
        ("诊断时间", sections['diagnosis_date']),
        ("诊断模式", sections['mode']),
        ("匹配政策数", f"{sections['total_policies']} 条"),
        ("平均综合匹配度", f"{sections['avg_score']} 分"),
    ]
    if sections['is_enhanced']:
        cover_info.append(("LLM 提供商", sections['llm_provider']))
        cover_info.append(("软条件评估数", f"{sections['soft_score_count']} 条"))

    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(12)

    doc.add_page_break()

    # ========== 执行摘要 ==========
    doc.add_heading('执行摘要', level=1)
    for line in sections['executive_summary'].split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('**') and stripped.endswith('**'):
            doc.add_paragraph(stripped.strip('*'))
        else:
            # 处理少量 markdown 加粗
            p = doc.add_paragraph()
            _add_markdown_text_to_docx_paragraph(p, stripped)

    # ========== 企业画像快照 ==========
    doc.add_heading('TOP3 推荐政策路线图', level=1)
    if sections['top3_policies']:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '排名'
        hdr_cells[1].text = '政策名称'
        hdr_cells[2].text = '诊断结果'
        hdr_cells[3].text = '综合分数'
        hdr_cells[4].text = '时间线建议'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        for p in sections['top3_policies']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(p['rank'])
            row_cells[1].text = p['policy_name']
            row_cells[2].text = p['diagnosis']
            row_cells[3].text = f"{p['combined_score']} 分"
            row_cells[4].text = p['timeline_advice']
            # 给诊断结果单元格上色
            _set_docx_cell_shading(row_cells[2], DIAGNOSIS_COLORS[p['diagnosis']]['hex'])

        for p in sections['top3_policies']:
            doc.add_heading(f"{p['rank']}. {p['policy_name']}", level=2)
            doc.add_paragraph(f"诊断结果：{p['diagnosis']}")
            doc.add_paragraph(f"综合分数：{p['combined_score']} 分")
            doc.add_paragraph(f"政策优先级：{p['priority']}")
            doc.add_paragraph(f"申报时间线：{p['timeline_advice']}")
            if p['deadline']:
                status = p['deadline_status']
                doc.add_paragraph(f"截止日：{p['deadline']}（{status['status_text']}）")
            doc.add_paragraph(f"扶持内容：{p['benefit']}")
            if p['key_gaps']:
                doc.add_paragraph(f"关键差距：{'、'.join(p['key_gaps'])}")
    else:
        doc.add_paragraph('暂无推荐政策。')

    # ========== 雷达图 / 维度分数 ==========
    doc.add_heading('企业能力雷达图与维度分数', level=1)
    if radar_image_bytes:
        image_stream = io.BytesIO(radar_image_bytes)
        doc.add_picture(image_stream, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        doc.add_paragraph('（未安装 kaleido，以下用表格展示维度分数）')
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = '能力维度'
        hdr[1].text = '分数'
        hdr[2].text = '评价'
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
        for dim, score in sections['capability_scores'].items():
            level = "强" if score >= 80 else "良" if score >= 60 else "弱"
            row = table.add_row().cells
            row[0].text = dim
            row[1].text = f"{score} 分"
            row[2].text = level

    doc.add_paragraph('维度短板建议：')
    for item in sections['capability_assessment']:
        doc.add_paragraph(item, style='List Bullet')

    # ========== 诊断结果总览 ==========
    doc.add_heading('全部政策诊断明细', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = '诊断结果'
    hdr[1].text = '数量'
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
    for diagnosis in ["立即申报", "培育申报", "持续关注", "暂不适合"]:
        row = table.add_row().cells
        row[0].text = diagnosis
        row[1].text = f"{sections['summary_counts'].get(diagnosis, 0)} 条"
        _set_docx_cell_shading(row[0], DIAGNOSIS_COLORS[diagnosis]['hex'])
    doc.add_paragraph(f"平均综合匹配度：{sections['avg_score']} 分")

    # ========== 详细诊断结果 ==========
    doc.add_heading('详细诊断结果', level=1)
    diagnosis_order = ["立即申报", "培育申报", "持续关注", "暂不适合"]
    for diagnosis in diagnosis_order:
        items = [r for r in sections['detailed_results'] if r['diagnosis'] == diagnosis]
        if not items:
            continue

        doc.add_heading(f"{diagnosis}（{len(items)} 条）", level=2)
        for r in items:
            doc.add_heading(r['policy_name'], level=3)

            if sections['is_enhanced'] and 'combined_score' in r:
                doc.add_paragraph(f"综合匹配度：{r['combined_score']} 分（硬 {r['hard_score']} + 软 {r.get('soft_score', 'N/A')}）")
            else:
                doc.add_paragraph(f"匹配度：{r['match_score']} 分")

            doc.add_paragraph(f"政策层级：{r['level']}")
            doc.add_paragraph(f"申报截止：{r['deadline']}")
            doc.add_paragraph(f"扶持内容：{r['benefit']}")
            doc.add_paragraph(f"诊断理由：{r['reason']}")

            if r['failed']:
                doc.add_paragraph('差距：')
                for item in r['failed']:
                    doc.add_paragraph(item, style='List Bullet 2')

            if r['unknown']:
                doc.add_paragraph('需补充数据：')
                for item in r['unknown']:
                    doc.add_paragraph(item, style='List Bullet 2')

            if sections['is_enhanced'] and r.get('soft_score') is not None:
                doc.add_paragraph(f"LLM 软条件评估：{r['soft_score']} 分（置信度：{r.get('confidence', '未知')}）")
                doc.add_paragraph(f"综合评估：{r.get('soft_assessment', '')}")

                if r.get('strengths'):
                    doc.add_paragraph('优势：')
                    for item in r['strengths']:
                        doc.add_paragraph(item, style='List Bullet 2')

                if r.get('weaknesses'):
                    doc.add_paragraph('短板：')
                    for item in r['weaknesses']:
                        doc.add_paragraph(item, style='List Bullet 2')

                if r.get('cultivation_suggestions'):
                    doc.add_paragraph('培育建议：')
                    for item in r['cultivation_suggestions']:
                        doc.add_paragraph(item, style='List Bullet 2')

            if r.get('material_outline'):
                outline = r['material_outline']
                doc.add_paragraph(f"申报可行性：{outline.get('applicability', '')}")
                doc.add_paragraph('申报材料大纲：')
                for section in outline.get('outline', []):
                    doc.add_paragraph(
                        f"{section.get('section', '')}：{'; '.join(section.get('content', []))}",
                        style='List Bullet 2'
                    )
                if outline.get('key_attachments'):
                    doc.add_paragraph('关键附件清单：')
                    for item in outline['key_attachments']:
                        doc.add_paragraph(item, style='List Bullet 2')
                if outline.get('gap_fill_plan'):
                    doc.add_paragraph('差距补齐计划：')
                    for item in outline['gap_fill_plan']:
                        doc.add_paragraph(item, style='List Bullet 2')
                if outline.get('notes'):
                    doc.add_paragraph(f"特别提醒：{outline['notes']}")

    # ========== 下一步行动建议 ==========
    doc.add_heading('下一步行动建议', level=1)
    doc.add_heading('立即行动（0-30 天）', level=2)
    for item in sections['next_steps']['immediate']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('中期培育（1-12 个月）', level=2)
    for item in sections['next_steps']['medium']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('长期关注（1-2 年）', level=2)
    for item in sections['next_steps']['long']:
        doc.add_paragraph(item, style='List Bullet')

    # 页脚说明
    doc.add_paragraph()
    footer = doc.add_paragraph('本报告由企业政策诊断辅导智能体自动生成')
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_markdown_text_to_docx_paragraph(paragraph, text: str):
    """将包含 **加粗** 的 markdown 文本写入 docx 段落"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part.strip('*'))
            run.bold = True
        else:
            paragraph.add_run(part)


def _set_docx_cell_shading(cell, hex_color: str):
    """设置 docx 单元格背景色"""
    from docx.oxml import parse_xml
    from docx.shared import RGBColor
    fill = hex_color.lstrip('#')
    shading_elm = parse_xml(
        '<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" '
        f'w:fill=\"{fill}\"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)
    # 设置文字颜色为白色，避免在深色背景上看不清（此处颜色都偏深，统一白字）
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)


def build_pdf_report(
    result: Dict[str, Any],
    capability_scores: Optional[Dict[str, int]] = None,
    radar_image_bytes: Optional[bytes] = None
) -> bytes:
    """
    生成 PDF 报告，返回 bytes

    参数：
        result: 诊断结果字典
        capability_scores: 企业能力维度分数（为 None 时从 data/enterprise.json 读取）
        radar_image_bytes: 雷达图图片 bytes（为 None 时使用文字表格 fallback）
    """
    from fpdf import FPDF

    font_path = _find_cjk_font()
    if not font_path:
        raise RuntimeError(
            "未找到中文字体，PDF 导出无法显示中文。"
            "请在 Windows 系统字体目录保留 simsun、simhei 或 msyh 字体之一。"
        )

    if capability_scores is None:
        enterprise_file = "data/enterprise.json"
        if os.path.exists(enterprise_file):
            with open(enterprise_file, 'r', encoding='utf-8') as f:
                enterprise = json.load(f)
            capability_scores = enterprise.get('capability_scores') or calculate_dimension_scores(enterprise)
        else:
            capability_scores = calculate_dimension_scores({})

    sections = build_report_sections(result, capability_scores)

    pdf = FPDF()
    pdf.add_font("cn", "", font_path, uni=True)
    pdf.add_font("cn", "B", font_path, uni=True)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    # ========== 封面 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 24)
    pdf.cell(0, 20, sections['title'], ln=True, align='C')
    pdf.ln(10)

    pdf.set_font("cn", "", 12)
    cover_info = [
        ("企业名称", sections['enterprise_name']),
        ("诊断时间", sections['diagnosis_date']),
        ("诊断模式", sections['mode']),
        ("匹配政策数", f"{sections['total_policies']} 条"),
        ("平均综合匹配度", f"{sections['avg_score']} 分"),
    ]
    if sections['is_enhanced']:
        cover_info.append(("LLM 提供商", sections['llm_provider']))
        cover_info.append(("软条件评估数", f"{sections['soft_score_count']} 条"))

    for label, value in cover_info:
        pdf.cell(0, 10, f"{label}：{value}", ln=True, align='C')

    # ========== 执行摘要 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "执行摘要", ln=True)
    pdf.ln(2)

    pdf.set_font("cn", "", 10)
    content_width = pdf.w - pdf.l_margin - pdf.r_margin

    for line in sections['executive_summary'].split('\n'):
        stripped = line.strip()
        if not stripped:
            pdf.ln(2)
            continue
        if stripped.startswith('### '):
            pdf.set_font("cn", "B", 13)
            pdf.cell(0, 8, stripped[4:], ln=True)
            pdf.set_font("cn", "", 10)
        elif stripped.startswith('## '):
            pdf.set_font("cn", "B", 14)
            pdf.cell(0, 8, stripped[3:], ln=True)
            pdf.set_font("cn", "", 10)
        elif stripped.startswith('- '):
            pdf.multi_cell(content_width, 5, '  - ' + stripped[2:])
        else:
            pdf.multi_cell(content_width, 5, stripped)

    # ========== 企业画像快照 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "企业画像快照", ln=True)
    pdf.ln(2)

    pdf.set_font("cn", "", 10)
    ep = sections['enterprise_profile']
    if ep['available']:
        pdf.multi_cell(content_width, 5, ep['text'])
        pdf.ln(3)
        profile_rows = [
            ('所属行业', f"{ep['industry']}（{ep['sub_industry']}）"),
            ('企业规模', f"{ep['scale']} / {ep['employees']} 人"),
            ('所在地区', ep['region']),
            ('成立年份', str(ep['founded_year'])),
            ('上年度营收', f"{ep['revenue']} 万元"),
            ('上年度利润', f"{ep['profit']} 万元"),
            ('研发投入', f"{ep['rd_investment']} 万元"),
            ('研发占比', str(ep['rd_ratio'])),
            ('研发人员', f"{ep['rd_team_size']} 人 / {ep['rd_team_ratio']}"),
            ('高新技术产品收入占比', str(ep['high_tech_income_ratio'])),
            ('知识产权', f"发明 {ep['invention_patents']} / 实用新型 {ep['utility_models']} / 软著 {ep['software_copyrights']} / 商标 {ep['trademarks']}"),
            ('已获资质', '、'.join(ep['qualifications']) if ep['qualifications'] else '—'),
            ('国家高新技术企业', '是' if ep['is_high_tech_enterprise'] else '否'),
            ('研发准备金制度', '已建立' if ep['rd_accounting_system'] else '未建立'),
            ('近三年重大事故', '有' if ep['has_major_accident'] else '无'),
        ]
        col_widths = [60, 100]
        pdf.set_font("cn", "B", 10)
        pdf.cell(col_widths[0], 7, '项目', border=1, align='C')
        pdf.cell(col_widths[1], 7, '内容', border=1, align='C')
        pdf.ln()
        pdf.set_font("cn", "", 10)
        for label, value in profile_rows:
            pdf.cell(col_widths[0], 7, label, border=1)
            pdf.cell(col_widths[1], 7, value, border=1)
            pdf.ln()
    else:
        pdf.multi_cell(content_width, 5, ep['text'])

    # ========== 诊断结果总览 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "诊断结果总览", ln=True)
    pdf.ln(2)

    col_widths = [80, 60]
    pdf.set_font("cn", "B", 10)
    pdf.cell(col_widths[0], 7, '诊断结果', border=1, align='C')
    pdf.cell(col_widths[1], 7, '数量', border=1, align='C')
    pdf.ln()

    pdf.set_font("cn", "", 10)
    for diagnosis in ["立即申报", "培育申报", "持续关注", "暂不适合"]:
        rgb = DIAGNOSIS_COLORS[diagnosis]['rgb']
        pdf.set_fill_color(*rgb)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(col_widths[0], 7, diagnosis, border=1, align='C', fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(col_widths[1], 7, f"{sections['summary_counts'].get(diagnosis, 0)} 条", border=1, align='C')
        pdf.ln()

    pdf.ln(3)
    pdf.set_font("cn", "", 10)
    metrics = sections['metrics']
    pdf.multi_cell(content_width, 5, f"平均综合匹配度：{sections['avg_score']} 分")
    pdf.multi_cell(content_width, 5, f"紧急截止政策数：{metrics['urgent_count']} 条")
    pdf.multi_cell(content_width, 5, f"已过期政策数：{metrics['expired_count']} 条")
    if metrics['nearest_deadline']:
        pdf.multi_cell(content_width, 5, f"最近截止日：{metrics['nearest_deadline']}（剩 {metrics['nearest_days']} 天）")

    # ========== 高频差距与短板分析 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "高频差距与短板分析", ln=True)
    pdf.ln(2)

    pdf.set_font("cn", "", 10)
    if sections['gap_analysis']:
        col_widths = [80, 60]
        pdf.set_font("cn", "B", 10)
        pdf.cell(col_widths[0], 7, '差距/缺失项', border=1, align='C')
        pdf.cell(col_widths[1], 7, '涉及政策数', border=1, align='C')
        pdf.ln()
        pdf.set_font("cn", "", 10)
        for name, count in sections['gap_analysis']:
            pdf.cell(col_widths[0], 7, name, border=1)
            pdf.cell(col_widths[1], 7, f"{count} 条", border=1, align='C')
            pdf.ln()
    else:
        pdf.multi_cell(content_width, 5, "未发现显著高频差距项。")

    # ========== TOP3 路线图 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "TOP3 推荐政策路线图", ln=True)
    pdf.ln(2)

    if sections['top3_policies']:
        # 表格
        col_widths = [15, 55, 30, 25, 55]
        row_height = 7
        pdf.set_font("cn", "B", 10)
        headers = ['排名', '政策名称', '诊断结果', '综合分数', '时间线建议']
        for i, h in enumerate(headers):
            pdf.cell(col_widths[i], row_height, h, border=1, align='C')
        pdf.ln()

        pdf.set_font("cn", "", 9)
        for p in sections['top3_policies']:
            # 诊断结果单元格使用填充色
            rgb = DIAGNOSIS_COLORS[p['diagnosis']]['rgb']
            pdf.set_fill_color(*rgb)
            pdf.set_text_color(255, 255, 255)

            pdf.cell(col_widths[0], row_height, str(p['rank']), border=1, align='C')
            pdf.cell(col_widths[1], row_height, p['policy_name'], border=1)
            pdf.cell(col_widths[2], row_height, p['diagnosis'], border=1, align='C', fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(col_widths[3], row_height, f"{p['combined_score']} 分", border=1, align='C')
            pdf.cell(col_widths[4], row_height, p['timeline_advice'], border=1)
            pdf.ln()

        pdf.ln(5)
        pdf.set_font("cn", "", 10)
        for p in sections['top3_policies']:
            pdf.set_font("cn", "B", 12)
            pdf.cell(0, 8, f"{p['rank']}. {p['policy_name']}", ln=True)
            pdf.set_font("cn", "", 10)
            pdf.multi_cell(content_width, 5, f"诊断结果：{p['diagnosis']}")
            pdf.multi_cell(content_width, 5, f"综合分数：{p['combined_score']} 分")
            pdf.multi_cell(content_width, 5, f"政策优先级：{p['priority']}")
            pdf.multi_cell(content_width, 5, f"申报时间线：{p['timeline_advice']}")
            if p['deadline']:
                status = p['deadline_status']
                pdf.multi_cell(content_width, 5, f"截止日：{p['deadline']}（{status['status_text']}）")
            pdf.multi_cell(content_width, 5, f"扶持内容：{p['benefit']}")
            if p['key_gaps']:
                pdf.multi_cell(content_width, 5, f"关键差距：{'、'.join(p['key_gaps'])}")
            pdf.ln(2)
    else:
        pdf.multi_cell(content_width, 5, "暂无推荐政策。")

    # ========== 雷达图 / 维度分数 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "企业能力雷达图与维度分数", ln=True)
    pdf.ln(2)

    if radar_image_bytes:
        image_stream = io.BytesIO(radar_image_bytes)
        pdf.image(image_stream, x=30, w=150)
    else:
        pdf.set_font("cn", "", 10)
        pdf.multi_cell(content_width, 5, "（未安装 kaleido，以下用表格展示维度分数）")
        pdf.ln(2)
        col_widths = [60, 40, 40]
        pdf.set_font("cn", "B", 10)
        for i, h in enumerate(['能力维度', '分数', '评价']):
            pdf.cell(col_widths[i], 7, h, border=1, align='C')
        pdf.ln()
        pdf.set_font("cn", "", 10)
        for dim, score in sections['capability_scores'].items():
            level = "强" if score >= 80 else "良" if score >= 60 else "弱"
            pdf.cell(col_widths[0], 7, dim, border=1)
            pdf.cell(col_widths[1], 7, f"{score} 分", border=1, align='C')
            pdf.cell(col_widths[2], 7, level, border=1, align='C')
            pdf.ln()

    pdf.ln(5)
    pdf.set_font("cn", "B", 12)
    pdf.cell(0, 8, "维度短板建议", ln=True)
    pdf.set_font("cn", "", 10)
    for item in sections['capability_assessment']:
        pdf.multi_cell(content_width, 5, '- ' + item)

    # ========== 全部政策诊断明细 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "全部政策诊断明细", ln=True)
    pdf.ln(2)

    col_widths = [80, 60]
    pdf.set_font("cn", "B", 10)
    pdf.cell(col_widths[0], 7, '诊断结果', border=1, align='C')
    pdf.cell(col_widths[1], 7, '数量', border=1, align='C')
    pdf.ln()

    pdf.set_font("cn", "", 10)
    for diagnosis in ["立即申报", "培育申报", "持续关注", "暂不适合"]:
        rgb = DIAGNOSIS_COLORS[diagnosis]['rgb']
        pdf.set_fill_color(*rgb)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(col_widths[0], 7, diagnosis, border=1, align='C', fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(col_widths[1], 7, f"{sections['summary_counts'].get(diagnosis, 0)} 条", border=1, align='C')
        pdf.ln()

    pdf.ln(3)
    pdf.set_font("cn", "", 10)
    pdf.multi_cell(content_width, 5, f"平均综合匹配度：{sections['avg_score']} 分")

    # ========== 详细诊断结果 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "详细诊断结果", ln=True)
    pdf.ln(2)

    diagnosis_order = ["立即申报", "培育申报", "持续关注", "暂不适合"]
    for diagnosis in diagnosis_order:
        items = [r for r in sections['detailed_results'] if r['diagnosis'] == diagnosis]
        if not items:
            continue

        pdf.set_font("cn", "B", 14)
        rgb = DIAGNOSIS_COLORS[diagnosis]['rgb']
        pdf.set_fill_color(*rgb)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 9, f" {diagnosis}（{len(items)} 条）", ln=True, fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

        for r in items:
            pdf.set_font("cn", "B", 11)
            pdf.cell(0, 7, r['policy_name'], ln=True)
            pdf.set_font("cn", "", 10)

            if sections['is_enhanced'] and 'combined_score' in r:
                pdf.multi_cell(content_width, 5, f"综合匹配度：{r['combined_score']} 分（硬 {r['hard_score']} + 软 {r.get('soft_score', 'N/A')}）")
            else:
                pdf.multi_cell(content_width, 5, f"匹配度：{r['match_score']} 分")

            pdf.multi_cell(content_width, 5, f"政策层级：{r['level']}")
            pdf.multi_cell(content_width, 5, f"申报截止：{r['deadline']}")
            pdf.multi_cell(content_width, 5, f"扶持内容：{r['benefit']}")
            pdf.multi_cell(content_width, 5, f"诊断理由：{r['reason']}")

            if r['failed']:
                pdf.multi_cell(content_width, 5, "差距：")
                for item in r['failed']:
                    pdf.multi_cell(content_width, 5, '  - ' + item)

            if r['unknown']:
                pdf.multi_cell(content_width, 5, "需补充数据：")
                for item in r['unknown']:
                    pdf.multi_cell(content_width, 5, '  - ' + item)

            if sections['is_enhanced'] and r.get('soft_score') is not None:
                pdf.multi_cell(content_width, 5, f"LLM 软条件评估：{r['soft_score']} 分（置信度：{r.get('confidence', '未知')}）")
                pdf.multi_cell(content_width, 5, f"综合评估：{r.get('soft_assessment', '')}")

                if r.get('strengths'):
                    pdf.multi_cell(content_width, 5, "优势：")
                    for item in r['strengths']:
                        pdf.multi_cell(content_width, 5, '  - ' + item)

                if r.get('weaknesses'):
                    pdf.multi_cell(content_width, 5, "短板：")
                    for item in r['weaknesses']:
                        pdf.multi_cell(content_width, 5, '  - ' + item)

                if r.get('cultivation_suggestions'):
                    pdf.multi_cell(content_width, 5, "培育建议：")
                    for item in r['cultivation_suggestions']:
                        pdf.multi_cell(content_width, 5, '  - ' + item)

            if r.get('material_outline'):
                outline = r['material_outline']
                pdf.multi_cell(content_width, 5, f"申报可行性：{outline.get('applicability', '')}")
                pdf.multi_cell(content_width, 5, "申报材料大纲：")
                for section in outline.get('outline', []):
                    pdf.multi_cell(content_width, 5, '  - ' + section.get('section', '') + '：' + '; '.join(section.get('content', [])))
                if outline.get('key_attachments'):
                    pdf.multi_cell(content_width, 5, "关键附件清单：")
                    for item in outline['key_attachments']:
                        pdf.multi_cell(content_width, 5, '  - ' + item)
                if outline.get('gap_fill_plan'):
                    pdf.multi_cell(content_width, 5, "差距补齐计划：")
                    for item in outline['gap_fill_plan']:
                        pdf.multi_cell(content_width, 5, '  - ' + item)
                if outline.get('notes'):
                    pdf.multi_cell(content_width, 5, f"特别提醒：{outline['notes']}")

            pdf.ln(2)

    # ========== 下一步行动建议 ==========
    pdf.add_page()
    pdf.set_font("cn", "B", 16)
    pdf.cell(0, 10, "下一步行动建议", ln=True)
    pdf.ln(2)

    sections_map = [
        ("立即行动（0-30 天）", sections['next_steps']['immediate']),
        ("中期培育（1-12 个月）", sections['next_steps']['medium']),
        ("长期关注（1-2 年）", sections['next_steps']['long']),
    ]
    for title, items in sections_map:
        pdf.set_font("cn", "B", 12)
        pdf.cell(0, 8, title, ln=True)
        pdf.set_font("cn", "", 10)
        for item in items:
            pdf.multi_cell(content_width, 5, '- ' + item)
        pdf.ln(2)

    # 页脚
    pdf.ln(5)
    pdf.set_font("cn", "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 8, "本报告由企业政策诊断辅导智能体自动生成", ln=True, align='C')

    return bytes(pdf.output(dest="S"))


def _md_inline_to_html(text: str) -> str:
    """简单处理 **加粗** 为 <strong>"""
    return re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)


def _simple_md_to_html(md_text: str) -> str:
    """把 build_executive_summary 生成的 markdown 转成简单 HTML"""
    html_lines = []
    in_list = False

    for raw in md_text.split('\n'):
        line = raw.strip()
        if not line:
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            continue

        if line.startswith('### '):
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            html_lines.append(f'<h3>{_md_inline_to_html(line[4:])}</h3>')
        elif line.startswith('## '):
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            html_lines.append(f'<h2>{_md_inline_to_html(line[3:])}</h2>')
        elif line.startswith('- '):
            if not in_list:
                html_lines.append('<ul>')
                in_list = True
            html_lines.append(f'<li>{_md_inline_to_html(line[2:])}</li>')
        else:
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            html_lines.append(f'<p>{_md_inline_to_html(line)}</p>')

    if in_list:
        html_lines.append('</ul>')
    return '\n'.join(html_lines)


def build_html_report(
    result: Dict[str, Any],
    capability_scores: Optional[Dict[str, int]] = None,
    fig_radar: Optional[Any] = None,
    fig_pie: Optional[Any] = None,
    fig_timeline: Optional[Any] = None,
    fig_capability: Optional[Any] = None,
    fig_gap: Optional[Any] = None,
) -> str:
    """
    生成可离线/在线打开的独立 HTML 诊断报告。

    图表使用 Plotly CDN 渲染；如果打开环境无网络，图表区域会显示空白。
    如需完全离线，可改用图片嵌入（需安装 kaleido）。
    """
    if capability_scores is None:
        enterprise_file = "data/enterprise.json"
        if os.path.exists(enterprise_file):
            with open(enterprise_file, 'r', encoding='utf-8') as f:
                enterprise = json.load(f)
            capability_scores = enterprise.get('capability_scores') or calculate_dimension_scores(enterprise)
        else:
            capability_scores = calculate_dimension_scores({})

    sections = build_report_sections(result, capability_scores)
    metrics = sections['metrics']
    top3 = sections['top3_policies']

    # 图表 HTML（仅加载一次 Plotly JS）
    chart_divs: Dict[str, str] = {}
    figures = {
        "radar": fig_radar,
        "pie": fig_pie,
        "timeline": fig_timeline,
        "capability": fig_capability,
        "gap": fig_gap,
    }
    first = True
    for key, fig in figures.items():
        if fig is None:
            chart_divs[key] = ""
            continue
        chart_divs[key] = fig.to_html(
            full_html=False,
            include_plotlyjs='cdn' if first else False,
            default_height='400px',
        )
        first = False

    # 企业画像表格行
    ep_rows = ""
    ep = sections['enterprise_profile']
    if ep['available']:
        profile_rows = [
            ("所属行业", f"{ep['industry']}（{ep['sub_industry']}）"),
            ("企业规模", f"{ep['scale']} / {ep['employees']} 人"),
            ("所在地区", ep['region']),
            ("成立年份", str(ep['founded_year'])),
            ("上年度营收", f"{ep['revenue']} 万元"),
            ("上年度利润", f"{ep['profit']} 万元"),
            ("研发投入", f"{ep['rd_investment']} 万元"),
            ("研发占比", str(ep['rd_ratio'])),
            ("研发人员", f"{ep['rd_team_size']} 人 / {ep['rd_team_ratio']}"),
            ("高新技术产品收入占比", str(ep['high_tech_income_ratio'])),
            ("知识产权", f"发明 {ep['invention_patents']} / 实用新型 {ep['utility_models']} / 软著 {ep['software_copyrights']} / 商标 {ep['trademarks']}"),
            ("已获资质", "、".join(ep['qualifications']) if ep['qualifications'] else "—"),
            ("国家高新技术企业", "是" if ep['is_high_tech_enterprise'] else "否"),
            ("研发准备金制度", "已建立" if ep['rd_accounting_system'] else "未建立"),
            ("近三年重大事故", "有" if ep['has_major_accident'] else "无"),
        ]
        ep_rows = "\n".join(
            f"        <tr><td><strong>{label}</strong></td><td>{value}</td></tr>"
            for label, value in profile_rows
        )

    # 诊断统计表格
    count_rows = "\n".join(
        f"        <tr><td>{diag}</td><td>{sections['summary_counts'].get(diag, 0)} 条</td></tr>"
        for diag in ["立即申报", "培育申报", "持续关注", "暂不适合"]
    )

    # TOP3 卡片
    top3_cards = ""
    for p in top3:
        deadline_text = f"截止：{p['deadline']}（{p['deadline_status']['status_text']}）" if p['deadline'] else "无固定截止日"
        key_gaps_html = ""
        if p['key_gaps']:
            key_gaps_html = f"      <p><strong>关键差距：</strong>{'、'.join(p['key_gaps'])}</p>\n"
        top3_cards += f"""
    <div class="top3-card top3-card-{p['diagnosis']}">
      <div class="top3-title">{p['rank']}. {p['policy_name']}</div>
      <div class="top3-meta">{p['diagnosis']} · 综合 {p['combined_score']} 分 · 政策优先级 {p['priority']}</div>
      <p><strong>时间线：</strong>{p['timeline_advice']}</p>
      <div class="top3-meta">{deadline_text}</div>
      <p><strong>扶持内容：</strong>{p['benefit']}</p>
      {key_gaps_html}
    </div>"""

    # 维度分数表格
    score_rows = "\n".join(
        f"        <tr><td>{dim}</td><td>{score}</td><td>{'强' if score >= 80 else '良' if score >= 60 else '弱'}</td></tr>"
        for dim, score in capability_scores.items()
    )

    # 差距分析
    gap_rows = "\n".join(
        f"        <tr><td>{name}</td><td>{count} 条</td></tr>"
        for name, count in sections['gap_analysis']
    ) if sections['gap_analysis'] else "        <tr><td colspan='2'>未发现显著高频差距项</td></tr>"

    # 行动建议
    action_cards = "\n".join(
        f"""
    <div class="action-card">
      <div class="action-title">{icon} {title}</div>
      <ul>{''.join(f'<li>{item}</li>' for item in items)}</ul>
    </div>"""
        for title, items, icon in [
            ("立即行动（0-30 天）", sections['next_steps']['immediate'], "⚡"),
            ("中期培育（1-12 个月）", sections['next_steps']['medium'], "🌱"),
            ("长期关注（1-2 年）", sections['next_steps']['long'], "🔭"),
        ]
    )

    nearest_deadline_html = ""
    if metrics['nearest_deadline']:
        nearest_deadline_html = f"      <li><strong>最近截止日：</strong>{metrics['nearest_deadline']}（剩 {metrics['nearest_days']} 天）</li>\n"

    html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{sections['enterprise_name']} - 企业政策诊断报告</title>
  <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
  <style>
    :root {{
      --bg: #f5f5f7;
      --card: #ffffff;
      --text: #1d1d1f;
      --muted: #6e6e73;
      --border: rgba(0,0,0,0.08);
      --blue: #0071e3;
      --blue-light: rgba(0,113,227,0.12);
      --green: #34c759;
      --green-light: rgba(52,199,89,0.12);
      --orange: #ff9500;
      --orange-light: rgba(255,149,0,0.12);
      --red: #ff3b30;
      --red-light: rgba(255,59,48,0.12);
    }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
      background: var(--bg);
      color: var(--text);
      line-height: 1.6;
      margin: 0;
      padding: 2rem 1rem;
    }}
    .container {{
      max-width: 960px;
      margin: 0 auto;
    }}
    .header {{
      text-align: center;
      margin-bottom: 2rem;
    }}
    .header h1 {{
      font-size: 2rem;
      margin: 0 0 0.5rem;
      letter-spacing: -0.02em;
    }}
    .header .meta {{
      color: var(--muted);
      font-size: 0.95rem;
    }}
    .card {{
      background: var(--card);
      border: 1px solid var(--border);
      border-radius: 20px;
      padding: 1.5rem;
      margin-bottom: 1.25rem;
      box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    }}
    .card h2 {{
      font-size: 1.15rem;
      margin: 0 0 1rem;
      padding-bottom: 0.65rem;
      border-bottom: 1px solid var(--border);
    }}
    .metrics {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
      gap: 1rem;
    }}
    .metric {{
      background: var(--card);
      border: 1px solid var(--border);
      border-radius: 16px;
      padding: 1rem;
      text-align: center;
      box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    }}
    .metric .value {{
      font-size: 1.6rem;
      font-weight: 700;
    }}
    .metric .label {{
      font-size: 0.85rem;
      color: var(--muted);
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin-top: 0.5rem;
    }}
    th, td {{
      padding: 0.6rem 0.75rem;
      border-bottom: 1px solid var(--border);
      text-align: left;
    }}
    th {{
      font-weight: 600;
      background: rgba(0,0,0,0.02);
    }}
    .top3-card {{
      border-radius: 12px;
      padding: 1.1rem 1.25rem;
      margin-bottom: 0.85rem;
      border: 1px solid var(--border);
      border-left: 5px solid #d1d1d6;
      background: var(--card);
    }}
    .top3-card-立即申报 {{ border-left-color: var(--green); background: var(--green-light); }}
    .top3-card-培育申报 {{ border-left-color: var(--orange); background: var(--orange-light); }}
    .top3-card-持续关注 {{ border-left-color: var(--blue); background: var(--blue-light); }}
    .top3-card-暂不适合 {{ border-left-color: var(--red); background: var(--red-light); }}
    .top3-title {{ font-weight: 700; font-size: 1.05rem; }}
    .top3-meta {{ color: var(--muted); font-size: 0.85rem; margin: 0.25rem 0; }}
    .grid-2 {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(320px, 1fr));
      gap: 1rem;
    }}
    .action-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(260px, 1fr));
      gap: 1rem;
    }}
    .action-card {{
      background: var(--card);
      border: 1px solid var(--border);
      border-radius: 16px;
      padding: 1.25rem;
    }}
    .action-title {{ font-weight: 700; margin-bottom: 0.75rem; }}
    .footer {{
      text-align: center;
      color: var(--muted);
      font-size: 0.85rem;
      margin-top: 2rem;
    }}
    ul {{ padding-left: 1.2rem; }}
    .chart {{ margin: 1rem 0; }}
  </style>
</head>
<body>
  <div class="container">
    <div class="header">
      <h1>📋 企业政策诊断报告</h1>
      <div class="meta">
        <strong>{sections['enterprise_name']}</strong> · 诊断时间：{sections['diagnosis_date']} · 匹配政策：{sections['total_policies']} 条
      </div>
    </div>

    <div class="metrics">
      <div class="metric"><div class="value">{sections['avg_score']}分</div><div class="label">平均匹配度</div></div>
      <div class="metric"><div class="value">{metrics['immediate']}条</div><div class="label">立即申报</div></div>
      <div class="metric"><div class="value">{metrics['cultivate']}条</div><div class="label">培育申报</div></div>
      <div class="metric"><div class="value">{metrics['urgent_count']}条</div><div class="label">紧急截止</div></div>
      <div class="metric"><div class="value">{metrics['watch']}条</div><div class="label">持续关注</div></div>
    </div>

    <div class="card">
      <h2>📄 执行摘要</h2>
      {_simple_md_to_html(sections['executive_summary'])}
    </div>

    <div class="card">
      <h2>🏢 企业画像快照</h2>
      {f'<p>{ep["text"]}</p>' if ep['available'] else f'<p>{ep["text"]}</p>'}
      <table>
        <tbody>
{ep_rows}
        </tbody>
      </table>
    </div>

    <div class="card">
      <h2>🔍 诊断结果总览</h2>
      <div class="grid-2">
        <div class="chart">{chart_divs['pie']}</div>
        <div>
          <table>
            <thead><tr><th>诊断结果</th><th>数量</th></tr></thead>
            <tbody>
{count_rows}
            </tbody>
          </table>
          <ul>
            <li><strong>平均综合匹配度：</strong>{sections['avg_score']} 分</li>
            <li><strong>紧急截止政策：</strong>{metrics['urgent_count']} 条</li>
            <li><strong>已过期政策：</strong>{metrics['expired_count']} 条</li>
{nearest_deadline_html}
          </ul>
        </div>
      </div>
    </div>

    <div class="card">
      <h2>🎯 TOP3 推荐政策路线图</h2>
      <div class="chart">{chart_divs['timeline']}</div>
      {top3_cards}
    </div>

    <div class="card">
      <h2>📈 企业能力与差距分析</h2>
      <div class="grid-2">
        <div class="chart">{chart_divs['radar']}</div>
        <div class="chart">{chart_divs['capability']}</div>
      </div>
      <div class="grid-2">
        <div>
          <table>
            <thead><tr><th>能力维度</th><th>分数</th><th>评价</th></tr></thead>
            <tbody>
{score_rows}
            </tbody>
          </table>
        </div>
        <div class="chart">{chart_divs['gap']}</div>
      </div>
      <h3>维度短板建议</h3>
      <ul>
{''.join(f'        <li>{item}</li>\n' for item in sections['capability_assessment'])}
      </ul>
    </div>

    <div class="card">
      <h2>🚀 下一步行动建议</h2>
      <div class="action-grid">
{action_cards}
      </div>
    </div>

    <div class="footer">本报告由企业政策诊断辅导智能体自动生成</div>
  </div>
</body>
</html>
"""
    return html

